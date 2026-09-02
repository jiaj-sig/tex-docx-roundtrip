"""
Apply text-level changes from an edited Word document back to a LaTeX source file.

This script performs best-effort synchronization:
  - It ignores headings, captions, and very short paragraphs.
  - It locates the corresponding LaTeX paragraph by rendering the .tex text
    the same way pandoc does (removing citations and math, unescaping special
    chars) and matching it against the original Word text.
  - It computes a word-level diff between original and edited Word text.
  - For each changed word block that appears verbatim in the original .tex
    paragraph, it replaces that block with the edited version.
  - It reports any changes it could not apply automatically.

Usage:
    python apply_text_changes.py <original.tex> <original.docx> <edited.docx> [output.tex]

If output.tex is omitted, the original .tex is overwritten after a backup is made.
"""
import argparse
import difflib
import os
import re
import shutil
import sys
from docx import Document


# Styles pandoc typically uses for body text
_BODY_STYLES = {"body text", "first paragraph", "normal", "abstract"}
# Styles to skip because structure changes here are risky
_SKIP_STYLES = {
    "title", "subtitle", "heading 1", "heading 2", "heading 3", "heading 4",
    "image caption", "captioned figure", "table", "table caption", "author",
    "abstract title", "reference", "bibliography",
}


def is_body_paragraph(style_name, text):
    s = style_name.lower()
    if s in _SKIP_STYLES:
        return False
    if s in _BODY_STYLES:
        return True
    if text.strip().startswith(("Figure ", "Table ", "Fig. ", "Section ", "\\section")):
        return False
    return True


def normalize_text(text):
    """Normalize whitespace variants to a single regular space."""
    for ch in ("\xa0", "\u2002", "\u2003", "\u2009", "\u202f"):
        text = text.replace(ch, " ")
    text = text.replace("\u00ad", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def render_tex_for_matching(tex):
    """Create a plain-text rendering of a LaTeX snippet matching pandoc output."""
    s = tex
    # Remove citation commands
    s = re.sub(r"\\cite\{[^}]*\}", "", s)
    # Non-breaking spaces and ties become regular spaces
    s = s.replace("~", " ")
    # LaTeX en-dash/em-dash
    s = s.replace("---", "—")
    s = s.replace("--", "–")
    # Remove inline math entirely
    s = re.sub(r"\$[^$]+\$", "", s)
    # Unescape common LaTeX special characters that pandoc renders literally
    for escaped, plain in [
        ("\\%", "%"),
        ("\\&", "&"),
        ("\\#", "#"),
        ("\\$", "$"),
        ("\\_", "_"),
        ("\\{", "{"),
        ("\\}", "}"),
        ("\\~", "~"),
        ("\\^", "^"),
    ]:
        s = s.replace(escaped, plain)
    # Normalize whitespace
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def find_tex_span(tex, old_word_text):
    """Find (start, end) in tex that corresponds to old_word_text.

    Returns the span of the best unique fuzzy match, or None.
    """
    rendered_full = render_tex_for_matching(tex)
    old_norm = normalize_text(old_word_text)
    words = [w for w in re.split(r"\s+", old_norm) if w]
    if not words:
        return None
    escaped = [re.escape(w) for w in words]
    pattern = r"[ \t\n\r~]*".join(escaped)
    matches = list(re.finditer(pattern, rendered_full, re.DOTALL))
    if len(matches) != 1:
        return None
    rm = matches[0]

    # Map rendered position back to original tex position.
    rend_pos = 0
    tex_pos = 0
    start = None
    end = None
    while tex_pos < len(tex) and rend_pos <= rm.end():
        if rend_pos == rm.start():
            start = tex_pos
        ch = tex[tex_pos]
        # Citation command: skip it entirely
        if ch == "\\" and tex[tex_pos:tex_pos + 5] == "\\cite":
            brace = tex.find("}", tex_pos)
            tex_pos = brace + 1 if brace != -1 else tex_pos + 1
            continue
        # Unescaped special chars contribute one rendered char
        if tex[tex_pos:tex_pos + 2] in ("\\%", "\\&", "\\#", "\\$", "\\_", "\\{", "\\}", "\\~", "\\^"):
            rend_pos += 1
            tex_pos += 2
            continue
        if ch == "~":
            rend_pos += 1
            tex_pos += 1
            continue
        if tex[tex_pos:tex_pos + 3] == "---":
            rend_pos += 1
            tex_pos += 3
            continue
        if tex[tex_pos:tex_pos + 2] == "--":
            rend_pos += 1
            tex_pos += 2
            continue
        if ch == "$":
            end_math = tex.find("$", tex_pos + 1)
            if end_math == -1:
                end_math = len(tex)
            tex_pos = end_math + 1
            continue
        if ch.isspace():
            rend_pos += 1
            tex_pos += 1
            continue
        rend_pos += 1
        tex_pos += 1
        if rend_pos == rm.end():
            end = tex_pos
            break
    if start is None or end is None:
        return None
    return start, end


def apply_word_changes(tex_span, old_text, new_text):
    """Apply word-level changes from old_text -> new_text onto tex_span."""
    old_norm = normalize_text(old_text)
    new_norm = normalize_text(new_text)
    if old_norm == new_norm:
        return tex_span

    old_words = old_norm.split()
    new_words = new_norm.split()
    sm = difflib.SequenceMatcher(None, old_words, new_words)
    result = tex_span
    for tag, i1, i2, j1, j2 in reversed(sm.get_opcodes()):
        if tag == "equal":
            continue
        old_phrase_words = old_words[i1:i2]
        new_phrase_words = new_words[j1:j2]
        old_phrase = " ".join(old_phrase_words)
        new_phrase = " ".join(new_phrase_words)
        if not old_phrase:
            continue
        idx = result.find(old_phrase)
        if idx == -1:
            pat = r"[ \t\n\r]*".join(re.escape(w) for w in old_phrase_words)
            m = re.search(pat, result)
            if not m:
                continue
            idx, end_idx = m.start(), m.end()
            result = result[:idx] + new_phrase + result[end_idx:]
        else:
            result = result[:idx] + new_phrase + result[idx + len(old_phrase):]
    return result


def apply_changes(tex_path, orig_docx, edit_docx, out_path):
    with open(tex_path, "r", encoding="utf-8") as f:
        tex = f.read()

    orig_doc = Document(orig_docx)
    edit_doc = Document(edit_docx)

    applied = []
    skipped = []

    orig_items = [(p.style.name if p.style else "None", p.text) for p in orig_doc.paragraphs]
    edit_items = [(p.style.name if p.style else "None", p.text) for p in edit_doc.paragraphs]

    n = max(len(orig_items), len(edit_items))
    for i in range(n):
        o_style, o_text = orig_items[i] if i < len(orig_items) else ("", "")
        e_style, e_text = edit_items[i] if i < len(edit_items) else ("", "")

        if o_text == e_text:
            continue
        if not is_body_paragraph(e_style or o_style, e_text or o_text):
            skipped.append({"index": i, "reason": "non-body paragraph", "old": o_text, "new": e_text})
            continue
        if len(o_text.strip()) < 10:
            skipped.append({"index": i, "reason": "original text too short", "old": o_text, "new": e_text})
            continue

        span = find_tex_span(tex, o_text)
        if span is None:
            skipped.append({"index": i, "reason": "no unique fuzzy match in .tex", "old": o_text, "new": e_text})
            continue

        start, end = span
        tex_span = tex[start:end]
        new_tex_span = apply_word_changes(tex_span, o_text, e_text)
        if new_tex_span == tex_span:
            skipped.append({"index": i, "reason": "could not apply word changes", "old": o_text, "new": e_text})
            continue
        tex = tex[:start] + new_tex_span + tex[end:]
        applied.append({"index": i, "old": o_text, "new": e_text})

    backup_path = tex_path + ".bak"
    if os.path.exists(backup_path):
        os.remove(backup_path)
    shutil.copy2(tex_path, backup_path)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(tex)

    return applied, skipped


def main():
    parser = argparse.ArgumentParser(description="Apply Word edits back to LaTeX")
    parser.add_argument("tex", help="Original .tex file")
    parser.add_argument("original_docx", help="Original .docx from tex2docx.py")
    parser.add_argument("edited_docx", help="User-edited .docx")
    parser.add_argument("output", nargs="?", help="Output .tex file (default: overwrite input)")
    args = parser.parse_args()

    out_path = os.path.abspath(args.output) if args.output else os.path.abspath(args.tex)
    applied, skipped = apply_changes(args.tex, args.original_docx, args.edited_docx, out_path)

    print(f"Applied {len(applied)} change(s)")
    if applied:
        for c in applied:
            print(f"  [#{c['index']}] {c['new'][:60]}")
    print(f"Skipped {len(skipped)} change(s)")
    if skipped:
        for c in skipped:
            print(f"  [#{c['index']}] {c['reason']}: {c['new'][:60]}")

    if skipped:
        print("\nPlease review skipped changes and apply them manually.", file=sys.stderr)
        sys.exit(2 if not applied else 0)


if __name__ == "__main__":
    main()
