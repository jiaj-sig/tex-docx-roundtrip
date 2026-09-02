"""
Compare an original pandoc-generated Word document with an edited version
and output a structured list of changed paragraphs.

Usage:
    python compare_docx.py <original.docx> <edited.docx> [diff.json]

Output JSON structure:
    [
      {
        "index": 0,
        "style": "Body Text",
        "old": "original paragraph text",
        "new": "edited paragraph text",
        "kind": "changed" | "added" | "removed"
      },
      ...
    ]
"""
import argparse
import json
import sys
from docx import Document


def iter_paragraphs(doc):
    """Yield (style_name, text) for each paragraph."""
    for p in doc.paragraphs:
        style = p.style.name if p.style else "None"
        text = p.text
        yield style, text


def align_and_diff(orig_items, edit_items):
    """Simple sequence alignment by index; flags additions/removals at the end."""
    diffs = []
    n = max(len(orig_items), len(edit_items))
    for i in range(n):
        o = orig_items[i] if i < len(orig_items) else ("", "")
        e = edit_items[i] if i < len(edit_items) else ("", "")
        if o[1] == e[1]:
            continue
        kind = "changed"
        if i >= len(orig_items):
            kind = "added"
        elif i >= len(edit_items):
            kind = "removed"
        diffs.append({
            "index": i,
            "style": e[0] if e[0] else o[0],
            "old": o[1],
            "new": e[1],
            "kind": kind,
        })
    return diffs


def main():
    parser = argparse.ArgumentParser(description="Compare two Word documents")
    parser.add_argument("original", help="Original .docx from tex2docx.py")
    parser.add_argument("edited", help="User-edited .docx")
    parser.add_argument("output", nargs="?", help="Output JSON file (default: stdout)")
    args = parser.parse_args()

    orig_doc = Document(args.original)
    edit_doc = Document(args.edited)

    orig_items = list(iter_paragraphs(orig_doc))
    edit_items = list(iter_paragraphs(edit_doc))

    diffs = align_and_diff(orig_items, edit_items)

    payload = {
        "original_paragraphs": len(orig_items),
        "edited_paragraphs": len(edit_items),
        "changes": diffs,
    }

    out = json.dumps(payload, ensure_ascii=False, indent=2)
    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(out)
        print(f"Wrote {len(diffs)} change(s) to {args.output}")
    else:
        print(out)


if __name__ == "__main__":
    main()
